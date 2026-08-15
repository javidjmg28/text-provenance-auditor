from __future__ import annotations

from pathlib import Path


def extract_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in {".txt", ".md", ".rst", ".csv", ".json", ".yaml", ".yml"}:
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "DOCX support requires the 'documents' extra: pip install -e '.[documents]'"
            ) from exc
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError(
                "PDF support requires the 'documents' extra: pip install -e '.[documents]'"
            ) from exc
        document = fitz.open(path)
        return "\n".join(page.get_text("text") for page in document)

    raise ValueError(f"Unsupported file type: {suffix or '<none>'}")
